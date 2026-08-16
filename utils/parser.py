import pdfplumber
from docx import Document
import io


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """
    Extract text content from a PDF file using multiple strategies.
    """
    text_parts = []
    
    # Strategy 1: pdfplumber (best for most PDFs)
    try:
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                # Try standard extraction
                page_text = page.extract_text()
                if page_text and page_text.strip():
                    text_parts.append(page_text.strip())
                    continue
                
                # Try with different tolerance
                page_text = page.extract_text(x_tolerance=3, y_tolerance=3)
                if page_text and page_text.strip():
                    text_parts.append(page_text.strip())
                    continue

                # Try tables
                tables = page.extract_tables()
                for table in tables:
                    for row in table:
                        row_text = " | ".join(str(cell).strip() for cell in row if cell)
                        if row_text.strip():
                            text_parts.append(row_text)
                
                # Try individual words
                words = page.extract_words()
                if words:
                    page_text = " ".join(w.get("text", "") for w in words)
                    if page_text.strip():
                        text_parts.append(page_text.strip())
    except Exception as e:
        print(f"[Parser Error] pdfplumber strategy failed: {e}")

    if text_parts:
        return "\n".join(text_parts).strip()

    # Strategy 2: pypdfium2 (fallback, handles some PDFs better)
    try:
        import pypdfium2 as pdfium
        pdf = pdfium.PdfDocument(file_bytes)
        for i in range(len(pdf)):
            page = pdf[i]
            textpage = page.get_textpage()
            page_text = textpage.get_text_range()
            if page_text and page_text.strip():
                text_parts.append(page_text.strip())
            textpage.close()
            page.close()
        pdf.close()
    except Exception as e:
        print(f"[Parser Error] pypdfium2 strategy failed: {e}")

    if text_parts:
        return "\n".join(text_parts).strip()

    # Strategy 3: pdfminer directly (another fallback)
    try:
        from pdfminer.high_level import extract_text as pdfminer_extract
        text = pdfminer_extract(io.BytesIO(file_bytes))
        if text and text.strip():
            return text.strip()
    except Exception as e:
        print(f"[Parser Error] pdfminer strategy failed: {e}")

    # Strategy 4: Gemini Multimodal Fallback (the ultimate backup for scanned/image PDFs)
    try:
        import os
        import google.generativeai as genai
        api_key = os.getenv("GEMINI_API_KEY")
        if api_key:
            genai.configure(api_key=api_key)
            print("[Parser] Attempting Gemini Multimodal PDF text extraction...")
            model = genai.GenerativeModel("gemini-2.5-flash")
            response = model.generate_content([
                {'mime_type': 'application/pdf', 'data': file_bytes},
                "Extract all the readable text from this resume PDF. Output the exact text as it is. Do not summarize or add metadata."
            ])
            if response.text and response.text.strip():
                print(f"[Parser] Gemini successfully extracted {len(response.text)} characters.")
                return response.text.strip()
    except Exception as e:
        print(f"[Parser Error] Gemini multimodal strategy failed: {e}")

    return ""


def extract_text_from_docx(file_bytes: bytes) -> str:
    """
    Extract text content from a DOCX file.
    """
    try:
        doc = Document(io.BytesIO(file_bytes))
    except Exception as e:
        raise ValueError(f"Error reading DOCX file: {str(e)}")
    
    text_parts = []
    
    # Extract from paragraphs
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            text_parts.append(paragraph.text.strip())
    
    # Extract from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                text_parts.append(row_text)

    # Extract from headers
    for section in doc.sections:
        header = section.header
        if header:
            for paragraph in header.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text.strip())

    return "\n".join(text_parts).strip()


def extract_text(filename: str, file_bytes: bytes) -> str:
    """
    Extract text from a file based on its extension.
    """
    lower_name = filename.lower()
    if lower_name.endswith('.pdf'):
        return extract_text_from_pdf(file_bytes)
    elif lower_name.endswith('.docx'):
        return extract_text_from_docx(file_bytes)
    else:
        raise ValueError(f"Unsupported file type: {filename}. Please upload a PDF or DOCX file.")
